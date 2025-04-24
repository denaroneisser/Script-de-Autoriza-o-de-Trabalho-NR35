import pandas as pd
from docx import Document
import os
from docx.shared import Pt

# Configurações
excel_file = "Participantes.xlsx"  # Nome do arquivo Excel
word_template = "MODELO2025.docx"  # Nome do arquivo Word modelo
output_folder = "documentos_gerados"  # Pasta de saída

# Criar pasta de saída se não existir
if not os.path.exists(output_folder):
    os.makedirs(output_folder)
    print(f"Pasta {output_folder} criada com sucesso!")

# Carregar dados do Excel
df = pd.read_excel(excel_file, dtype={"CPF": str})  # Garantir que o CPF seja tratado como string
print("Dados carregados:")
print(df.head())  # Exibe as primeiras linhas

# Função para substituir texto preservando formatação
def substituir_texto_completo(paragraph, substituicoes):
    full_text = "".join(run.text for run in paragraph.runs)
    for chave, valor in substituicoes.items():
        full_text = full_text.replace(chave, valor)  # Substitui todos os placeholders no texto completo
    
    if full_text != "".join(run.text for run in paragraph.runs):  # Só atualiza se houve substituição
        paragraph.clear()  # Remove runs antigos
        paragraph.add_run(full_text)  # Adiciona novo texto mantendo a formatação geral do parágrafo

# Função para ajustar a fonte de todo o documento para tamanho 13
def ajustar_tamanho_fonte(doc, tamanho=13):
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(tamanho)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(tamanho)

# Gerar documentos
for index, row in df.iterrows():
    try:
        print(f"Processando: {row['Nome']} - {row['CPF']}")
        doc = Document(word_template)
        print("Modelo carregado com sucesso!")

        # Formatar a data corretamente
        if isinstance(row["Data"], pd.Timestamp):
            data_formatada = row["Data"].strftime("%d/%m/%Y")
        else:
            data_formatada = str(row["Data"])

        # Criar dicionário de substituições
        substituicoes = {
            "{Nome}": str(row["Nome"]),
            "{nMatricula}": str(row["nMatricula"]),
            "{Funcao}": str(row["Funcao"]),
            "{Data}": data_formatada,
            "{CPF}": str(row["CPF"]),
        }

        # Substituir texto nos parágrafos
        for p in doc.paragraphs:
            substituir_texto_completo(p, substituicoes)

        # Substituir texto dentro das tabelas
        for table in doc.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for p in cell.paragraphs:
                        substituir_texto_completo(p, substituicoes)

        # Ajustar a fonte do documento para tamanho 13
        ajustar_tamanho_fonte(doc, tamanho=13)

        # Nome do arquivo de saída
        output_file = os.path.join(output_folder, f"{row['Nome']}_Autorização de Trabalho.docx")
        print(f"Arquivo de saída: {output_file}")

        # Salvar documento gerado
        doc.save(output_file)
        print(f"Documento gerado: {output_file}")

    except Exception as e:
        print(f"Erro ao processar {row['Nome']} - {row['CPF']}: {e}")

print("Todos os documentos foram gerados com sucesso!")
